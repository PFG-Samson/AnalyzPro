import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_user_manual(input_md_path, output_docx_path):
    """
    Reads a markdown file and converts it to a formatted Word document.
    """
    if not os.path.exists(input_md_path):
        print(f"Error: Input file not found at {input_md_path}")
        return

    document = Document()
    
    # Set up styles (basic)
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    with open(input_md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_block_content = []

    for line in lines:
        line = line.rstrip()
        
        # Handle Code Blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                in_code_block = False
                p = document.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                runner = p.add_run('\n'.join(code_block_content))
                runner.font.name = 'Courier New'
                runner.font.size = Pt(9)
                code_block_content = []
            else:
                # Start of code block
                in_code_block = True
            continue
        
        if in_code_block:
            code_block_content.append(line)
            continue

        # Headers
        if line.startswith('# '):
            document.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            document.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            document.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            document.add_heading(line[5:], level=3)
        
        # Horizontal Rule
        elif line.startswith('---'):
            document.add_page_break()
            
        # List Items (Basic support)
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            p = document.add_paragraph(line.strip()[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', line.strip()):
             # Remove the number and dot, let Word handle numbering or just treat as text for now
             # For simplicity in this script, we'll just add it as a paragraph, 
             # or we could try to use 'List Number' style but restarting numbering is tricky.
             # Let's stick to simple paragraph with the number for now to preserve original text.
             p = document.add_paragraph(line.strip(), style='List Number')

        # Tables (Basic text representation for now, parsing MD tables to Docx tables is complex)
        elif line.startswith('|'):
             # Just add as monospace text for now to preserve alignment roughly
             p = document.add_paragraph()
             runner = p.add_run(line)
             runner.font.name = 'Courier New'
             runner.font.size = Pt(9)

        # Normal Text
        else:
            if line.strip():
                document.add_paragraph(line)

    document.save(output_docx_path)
    print(f"Successfully created {output_docx_path}")

if __name__ == "__main__":
    input_file = "COMPLETE_DOCUMENTATION.md"
    output_file = "Analyz_User_Manual.docx"
    create_user_manual(input_file, output_file)
